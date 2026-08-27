"""
Sample Railway Document Generator.
Generates sample PDF and DOCX documents representing realistic multi-department
maintenance circulars, structured tables, and prose memos for runtime upload testing.
"""
import os
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "sample_documents")
os.makedirs(OUT_DIR, exist_ok=True)


def generate_docx_circular():
    """Generates a realistic multi-department weekly maintenance circular in DOCX format."""
    doc = Document()
    
    # Title
    title = doc.add_heading("NORTHERN RAILWAY - DIVISIONAL OPERATING CIRCULAR", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("REF: NR/OPTG/BLOCK-REQ/2026/W35 | DATE: 27-AUG-2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("1. Engineering & Track Maintenance Work Proposals", level=1)
    doc.add_paragraph(
        "The following track tamping, rail grinding, and ballast screening blocks are requested by Sr.DEN/Civil "
        "for urgent track stabilization and speed restriction removal on high-density corridors:"
    )

    # Structured Table 1: Civil & Electrical Joint Proposals
    table_data = [
        ["Job ID", "Dept", "Corridor", "KM Span", "Nature of Work", "Duration", "Permitted Window (IST)", "Machinery Req", "Priority"],
        ["REQ-NR-001", "Engineering", "NDLS-GZB", "KM 12.5 to 20.0", "Continuous Track Tamping (CSM)", "180 min", "01:00 to 04:30", "Track Tamper TTM-401", "P1"],
        ["REQ-NR-002", "Electrical", "NDLS-GZB", "KM 14.0 to 22.0", "OHE 25kV Catenary Periodic Overhaul", "150 min", "01:30 to 04:30", "Tower Wagon TW-3", "P2"],
        ["REQ-NR-003", "S&T", "NDLS-GZB", "KM 16.0 to 18.5", "Point Machine Replacement & Testing", "90 min", "02:00 to 04:00", "S&T Testing Team", "P3"],
        ["REQ-NR-004", "Engineering", "NDLS-CNB", "KM 45.0 to 58.0", "Rail Grinding (Speno Train)", "240 min", "00:30 to 05:00", "Speno Rail Grinder", "P1"],
        ["REQ-NR-005", "Electrical", "NDLS-CNB", "KM 48.0 to 60.0", "Insulator Replacement & Power Isolation", "180 min", "01:00 to 04:30", "Tower Wagon TW-1", "P2"],
        ["REQ-NR-006", "Engineering", "AGC-JHS", "KM 110.0 to 125.0", "Deep Screening & Ballast Cleaning", "210 min", "01:00 to 05:00", "Ballast Cleaner BCM", "P2"],
        ["REQ-NR-007", "S&T", "AGC-JHS", "KM 112.0 to 120.0", "Axle Counter Dual-Sensors Tuning", "120 min", "01:30 to 04:00", "S&T Testing Team", "P3"],
    ]

    t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    t.style = 'Table Grid'
    for r_idx, row in enumerate(table_data):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.text = val
            if r_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_heading("2. Special Traffic & Train Precedence Guidelines", level=1)
    doc.add_paragraph(
        "Train 12004 Vande Bharat Express on NDLS-CNB departing 06:00 arriving 11:30 (KM 0.0 to 440.0) has highest precedence. "
        "No maintenance block shall encroach beyond 05:00 IST on the NDLS-CNB mainline."
    )

    doc.add_heading("3. Unscheduled Emergency Works (Semi-Structured Memo)", level=1)
    doc.add_paragraph(
        "Item E-91: Electrical traction department SSE/TRD requests emergency power block for broken dropper replacement "
        "on corridor HWH-KGP between KM 78.0 and 85.5 from 02:00 to 04:30 IST. Power block required. Priority 1."
    )
    doc.add_paragraph(
        "Item E-92: P-Way gang requests 2.5 hour fishplate tightening on corridor HWH-KGP between KM 80.0 and 88.0 from 02:15 to 04:45 IST. Priority 2."
    )

    docx_path = os.path.join(OUT_DIR, "Northern_Railway_Maintenance_Circular.docx")
    doc.save(docx_path)
    print(f"Generated sample DOCX: {docx_path}")


def generate_pdf_requests():
    """Generates a PDF document with structured and prose maintenance requests."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors

        pdf_path = os.path.join(OUT_DIR, "Western_Railway_Block_Requisition.pdf")
        doc = SimpleDocTemplate(pdf_path, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        
        styles = getSampleStyleSheet()
        story = []

        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=16,
            leading=20,
            alignment=1,
            textColor=colors.HexColor('#0B132B')
        )
        story.append(Paragraph("WESTERN RAILWAY - DIVISIONAL BLOCK PROGRAMME", title_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("<b>Circular No:</b> WR/ENG/2026/AUG-W4 | <b>Target Date:</b> Tomorrow (IST Night Shift)", styles['Normal']))
        story.append(Spacer(1, 12))

        # Table
        table_data = [
            ["ID", "Dept", "Corridor", "KM Span", "Work Type", "Duration", "Window (IST)", "Machinery", "Priority"],
            ["WR-101", "Engineering", "MMCT-BVI", "KM 10.0 to 18.0", "Turnout Renewal & Tamping", "180 min", "01:00 - 04:30", "TTM Machine", "P1"],
            ["WR-102", "Electrical", "MMCT-BVI", "KM 12.0 to 20.0", "Catenary Wire Adjustment", "150 min", "01:30 - 04:30", "Tower Wagon", "P2"],
            ["WR-103", "S&T", "MMCT-BVI", "KM 11.0 to 16.0", "Signal Interlocking Verification", "120 min", "02:00 - 04:00", "Signal Team", "P3"],
            ["WR-104", "Engineering", "BVI-ST", "KM 95.0 to 110.0", "Track Grinding & Resurfacing", "210 min", "01:00 - 05:00", "Rail Grinder", "P2"],
            ["WR-105", "Electrical", "BVI-ST", "KM 98.0 to 112.0", "OHE Mast Inspection", "180 min", "01:30 - 04:30", "Tower Wagon", "P3"],
        ]

        t = Table(table_data, colWidths=[50, 65, 65, 75, 110, 50, 65, 60, 40])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1C2541')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8FAFC')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(t)
        story.append(Spacer(1, 15))

        # Incomplete sample request (Rule 8 test case)
        story.append(Paragraph("<b>Section 2: Pending Maintenance Notes (Flagged for Review)</b>", styles['Heading3']))
        story.append(Paragraph(
            "Note 44: Civil engineering crew requests 2 hour track inspection on corridor MMCT-BVI. (Missing KM span and time window - needs human planner completion).",
            styles['Normal']
        ))
        story.append(Spacer(1, 8))
        story.append(Paragraph(
            "Note 45: Train 12951 Mumbai Rajdhani on corridor BVI-ST departing 05:30 arriving 10:00 (KM 0 to 300).",
            styles['Normal']
        ))

        doc.build(story)
        print(f"Generated sample PDF: {pdf_path}")
    except Exception as e:
        print(f"PDF generation error: {e}")


if __name__ == "__main__":
    generate_docx_circular()
    try:
        generate_pdf_requests()
    except Exception as e:
        print("Note: reportlab not installed for PDF, DOCX generated.")