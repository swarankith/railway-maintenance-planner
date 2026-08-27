"""
Document Extractor for PDF and DOCX files.
Extracts raw text, structured tables, key-value blocks, and paragraphs from uploaded files.
"""
import io
import os
from typing import List, Dict, Any, Tuple


class DocumentContent:
    def __init__(self, filename: str):
        self.filename = filename
        self.raw_text: str = ""
        self.tables: List[List[List[str]]] = []  # List of tables, each table is rows of string cells
        self.sections: List[Dict[str, Any]] = []


def extract_from_pdf(file_bytes: bytes, filename: str) -> DocumentContent:
    """Extracts text and tables from PDF bytes using pdfplumber with pypdf fallback."""
    doc = DocumentContent(filename=filename)
    all_text_chunks = []

    # Attempt extraction via pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                all_text_chunks.append(page_text)
                
                # Extract tables
                page_tables = page.extract_tables()
                if page_tables:
                    for table in page_tables:
                        cleaned_table = []
                        for row in table:
                            cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                            if any(cleaned_row):
                                cleaned_table.append(cleaned_row)
                        if cleaned_table:
                            doc.tables.append(cleaned_table)
    except Exception as e:
        # Fallback to pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                txt = page.extract_text() or ""
                all_text_chunks.append(txt)
        except Exception as e2:
            all_text_chunks.append(f"[PDF Extraction Error: {str(e2)}]")

    doc.raw_text = "\n\n".join(all_text_chunks)
    return doc


def extract_from_docx(file_bytes: bytes, filename: str) -> DocumentContent:
    """Extracts text and tables from DOCX bytes using python-docx."""
    doc = DocumentContent(filename=filename)
    try:
        from docx import Document
        docx_doc = Document(io.BytesIO(file_bytes))
        
        # Extract paragraphs
        paragraphs = [p.text.strip() for p in docx_doc.paragraphs if p.text.strip()]
        doc.raw_text = "\n".join(paragraphs)

        # Extract tables
        for table in docx_doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data.append(row_data)
            if table_data:
                doc.tables.append(table_data)

    except Exception as e:
        doc.raw_text = f"[DOCX Extraction Error: {str(e)}]"

    return doc


def extract_document(file_bytes: bytes, filename: str) -> DocumentContent:
    """Dispatches extraction based on file extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_from_pdf(file_bytes, filename)
    elif ext in [".docx", ".doc"]:
        return extract_from_docx(file_bytes, filename)
    elif ext in [".txt", ".csv"]:
        content = file_bytes.decode("utf-8", errors="replace")
        doc = DocumentContent(filename=filename)
        doc.raw_text = content
        return doc
    else:
        # Generic attempt
        try:
            return extract_from_docx(file_bytes, filename)
        except:
            return extract_from_pdf(file_bytes, filename)
